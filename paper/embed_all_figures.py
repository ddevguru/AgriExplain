"""Embed B&W diagrams and colorful website screenshots into journal/conference papers."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph

PAPER_DIR = Path(__file__).resolve().parent
FIG = PAPER_DIR / "figures"
BW = FIG / "bw"
WEB = FIG / "website_color"

# (caption keywords, image path relative to FIG, colorful?)
RULES: list[tuple[list[str], str, bool]] = [
    # Colorful website UI screenshots
    (["landing page", "hero section", "next.js frontend"], "website_color/web_landing.png", True),
    (["login", "registration", "sign-in", "authentication page"], "website_color/web_login.png", True),
    (["camera panel", "thermal colormap", "plant diagnosis"], "website_color/web_farmer_dashboard.png", True),
    (["sensor cards", "yield risk badge", "shap panel", "advisories"], "website_color/web_farmer_dashboard.png", True),
    (["sensor reading cards grid", "full sensor reading grid"], "website_color/web_farmer_dashboard.png", True),
    (["shap explanation panel", "top-3 feature"], "website_color/web_farmer_dashboard.png", True),
    (["historical sensor time-series", "recharts"], "website_color/web_farmer_dashboard.png", True),
    (["admin dashboard", "multi-farm sensor health", "stale-data"], "website_color/web_admin_sensors.png", True),
    (["admin dashboard", "model performance comparison"], "website_color/web_admin_models.png", True),
    (["admin dashboard", "sensor health overview"], "website_color/web_admin_overview.png", True),
    (["capacitor", "android", "mobile device"], "website_color/web_farmer_dashboard.png", True),
    (["website"], "website_color/web_landing.png", True),
    # B&W technical diagrams
    (["methodology flowchart", "system design methodology"], "bw/methodology_flowchart.png", False),
    (["end-to-end", "system architecture"], "bw/architecture.png", False),
    (["deployment topology", "four-tier architecture", "multi-farm"], "bw/deployment_topology.png", False),
    (["hardware functional block", "gpio map", "hardware gpio"], "bw/hardware_block.png", False),
    (["entity-relationship", "relational schema", "six-table"], "bw/er_diagram.png", False),
    (["ingestion pipeline", "six-stage", "flowchart"], "bw/sensor_pipeline.png", False),
    (["feature vector", "nine-dimensional"], "bw/feature_vector.png", False),
    (["agrifusion", "fusion algorithm", "weighted probability fusion"], "bw/agrifusion_diagram.png", False),
    (["damping curve", "disagreement-aware"], "bw/damping_curve.png", False),
    (["model performance comparison", "weighted f1"], "bw/model_comparison.png", False),
    (["confusion matric"], "bw/confusion_matrices.png", False),
    (["shap", "tree shap", "feature attribution", "feature importance"], "bw/shap_chart.png", False),
    (["dataset", "per-crop sample", "yield-class distribution", "yield risk distribution", "yield risk level"], "bw/dataset_distribution.png", False),
    (["24-hour", "telemetry", "time-series"], "bw/telemetry_chart.png", False),
    (["npk reference", "crop npk", "modbus fallback"], "bw/npk_chart.png", False),
]


def _has_image_before(caption_para: Paragraph) -> bool:
    prev = caption_para._element.getprevious()
    if prev is None or not prev.tag.endswith("p"):
        return False
    return bool(prev.findall(".//" + qn("a:blip")))


def _clear_orphan_images(caption_para: Paragraph) -> None:
    """Remove empty image paragraphs directly before caption."""
    prev = caption_para._element.getprevious()
    while prev is not None and prev.tag.endswith("p"):
        has_text = bool("".join(prev.itertext()).strip())
        has_blip = bool(prev.findall(".//" + qn("a:blip")))
        if has_blip and not has_text:
            nxt = prev.getprevious()
            prev.getparent().remove(prev)
            prev = nxt
            continue
        break


def _resolve_image(caption: str) -> Path | None:
    cap = caption.lower()
    for keywords, rel, _ in RULES:
        if any(k in cap for k in keywords):
            path = FIG / rel
            if path.exists():
                return path
    # Fallback B&W by figure type
    if any(w in cap for w in ["dashboard", "login", "landing", "website", "farmer", "admin", "android"]):
        for name in ["web_landing.png", "web_login.png", "web_farmer_dashboard.png"]:
            p = WEB / name
            if p.exists():
                return p
    return None


def _insert_image_before(caption_para: Paragraph, image_path: Path, colorful: bool) -> None:
    _clear_orphan_images(caption_para)
    if _has_image_before(caption_para):
        prev = caption_para._element.getprevious()
        prev_para = Paragraph(prev, caption_para._parent)
        for run in list(prev_para.runs):
            if run._element.findall(".//" + qn("a:blip")):
                run._element.getparent().remove(run._element)
        prev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prev_para.add_run().add_picture(str(image_path), width=Inches(6.3))
        return

    img_para = caption_para.insert_paragraph_before()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(str(image_path), width=Inches(6.3))


def embed_figures(doc_path: Path, out_path: Path) -> int:
    doc = Document(str(doc_path))
    count = 0
    missing: list[str] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if not re.match(r"^Fig\.?\s*\d+", t, re.I):
            continue
        img = _resolve_image(t)
        if img is None:
            missing.append(t[:70])
            continue
        colorful = any(
            img == FIG / rel and c
            for keywords, rel, c in RULES
            if any(k in t.lower() for k in keywords)
        )
        _insert_image_before(para, img, colorful)
        count += 1

    doc.save(str(out_path))
    print(f"  embedded {count} figures -> {out_path.name}")
    if missing:
        print(f"  missing assets for {len(missing)} captions")
        for m in missing[:5]:
            print(f"    - {m}")
    return count


def main() -> None:
    jobs = [
        (
            PAPER_DIR / "AgriXplain_Journal_v3_SirReview.docx",
            PAPER_DIR / "AgriXplain_Journal_v3_Figures.docx",
        ),
        (
            PAPER_DIR / "AgriXplain_Conference_v3_SirReview.docx",
            PAPER_DIR / "AgriXplain_Conference_v3_Figures.docx",
        ),
    ]
    for src, dst in jobs:
        if not src.exists():
            print("SKIP", src)
            continue
        print("Processing", src.name)
        embed_figures(src, dst)


if __name__ == "__main__":
    main()
